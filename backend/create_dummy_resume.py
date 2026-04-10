import docx

doc = docx.Document()
doc.add_heading('Mock Dummy Resume', 0)
doc.add_paragraph('Name: Test User')
doc.add_paragraph('Skills: React, Node, Python, Django')
doc.add_paragraph('Experience: Software Engineer at Tech Corp. Developed web applications.')
doc.add_paragraph('Education: Computer Science, University of Examples')

doc.save('test_resume.docx')
print("Created test_resume.docx successfully.")
